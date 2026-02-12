"""Extract text from Word documents, chunked by heading."""

from pathlib import Path
from app.domain.documents import Chunk
from app.retrieval.processor.chunking_utils import split_text_to_chunks


def extract_docx(filepath: str) -> list[Chunk]:
    from docx import Document

    doc = Document(filepath)
    fname = Path(filepath).name
    chunks: list[Chunk] = []
    current_heading = "Introduction"
    current_text: list[str] = []
    seen_heading = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if current_text:
                section_content = "\n".join(current_text)
                for idx, chunk_text in enumerate(split_text_to_chunks(section_content), 1):
                    chunks.append(Chunk(
                        source_file=fname, source_type="docx",
                        section_title=f"{current_heading} — Part {idx}",
                        content=chunk_text,
                    ))
            current_heading = text
            seen_heading = True
            current_text = []
        else:
            current_text.append(text)

    if current_text:
        section_content = "\n".join(current_text)
        title = current_heading if seen_heading else "Full Document"
        for idx, chunk_text in enumerate(split_text_to_chunks(section_content), 1):
            chunks.append(Chunk(
                source_file=fname, source_type="docx",
                section_title=f"{title} — Part {idx}",
                content=chunk_text,
            ))

    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            chunks.append(Chunk(
                source_file=fname, source_type="docx",
                section_title=f"Table {i + 1}",
                content="\n".join(rows),
            ))

    return chunks
